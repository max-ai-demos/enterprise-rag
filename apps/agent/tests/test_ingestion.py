# apps/agent/tests/test_ingestion.py
import pytest
from pathlib import Path


def make_sample_pdf(path: Path):
    """Creates a test PDF using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "This is page one content about contracts and payment terms.")
    page.insert_text((72, 130), "The payment deadline is 30 days from invoice date.")
    doc.save(str(path))


def test_parse_pdf_returns_chunks_with_bbox(tmp_path):
    from app.rag.ingestion import parse_document
    pdf = tmp_path / "test.pdf"
    make_sample_pdf(pdf)
    chunks = parse_document(str(pdf), "pdf")
    assert len(chunks) > 0
    for chunk in chunks:
        assert "text" in chunk
        assert "page_num" in chunk
        assert "page_idx" in chunk
        assert "bbox" in chunk
        assert isinstance(chunk["bbox"], list) and len(chunk["bbox"]) == 4
        assert all(0 <= v <= 1000 for v in chunk["bbox"])
        assert chunk["page_num"] == 1
        assert chunk["page_idx"] == 1


def test_parse_txt_returns_chunks(tmp_path):
    from app.rag.ingestion import parse_document
    txt = tmp_path / "test.txt"
    txt.write_text("Line one.\nLine two.\nLine three." * 20)
    chunks = parse_document(str(txt), "txt")
    assert len(chunks) > 0
    assert all("text" in c for c in chunks)


def test_parse_docx_returns_chunks(tmp_path):
    from app.rag.ingestion import parse_document
    from docx import Document as DocxDocument
    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph about contracts.")
    doc.add_paragraph("Second paragraph about payments.")
    doc.save(str(docx_path))
    chunks = parse_document(str(docx_path), "docx")
    assert len(chunks) > 0
    assert all("paragraph_idx" in c for c in chunks)


def test_chunk_size_is_reasonable():
    from app.rag.ingestion import CHUNK_SIZE
    assert 300 <= CHUNK_SIZE <= 1000, f"CHUNK_SIZE {CHUNK_SIZE} out of expected range"


def test_parse_docx_includes_tables(tmp_path):
    """DOCX parsing should extract table rows as chunks."""
    from docx import Document as DocxDocument

    # Build a minimal DOCX with one paragraph and one table
    doc = DocxDocument()
    doc.add_paragraph("这是正文段落。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "产品名称"
    table.cell(0, 1).text = "单价"
    table.cell(1, 0).text = "笔记本电脑"
    table.cell(1, 1).text = "8000元"
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    from app.rag.ingestion import parse_document
    chunks = parse_document(str(docx_path), "docx")

    texts = [c["text"] for c in chunks]
    assert any("这是正文段落" in t for t in texts)
    assert any("产品名称" in t and "单价" in t for t in texts)
    assert any("笔记本电脑" in t and "8000元" in t for t in texts)
